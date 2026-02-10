#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para converter PDF de contrato em template DOCX com chaves substituíveis
"""

from pdf2docx import Converter
from docx import Document
import re
import sys

def converter_pdf_para_docx(pdf_path, docx_path):
    """Converte PDF para DOCX"""
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    print(f"✓ PDF convertido para DOCX: {docx_path}")

def substituir_dados_por_chaves(docx_path, output_path):
    """Substitui dados específicos por chaves de template"""
    doc = Document(docx_path)
    
    # Mapeamento de substituições
    substituicoes = {
        # Cliente
        'ALONSO PEREIRA PIMENTA': '{{cliente_nome}}',
        '052.853.621-48': '{{cliente_cpf_cnpj}}',
        '2196489 SEJUSP MS': '{{cliente_rg}}',
        'Rua Natal, nº 1310': '{{cliente_endereco}}',
        'Vila Cuiabá': '{{cliente_bairro}}',
        'Dourados': '{{cliente_cidade}}',
        'MS': '{{cliente_estado}}',
        '79841-010': '{{cliente_cep}}',
        
        # Empresa
        'MAB ENERGIA SOLAR': '{{empresa_razao_social}}',
        '38.068.450/0001-99': '{{empresa_cnpj}}',
        'Rua Maria Rosa, nº S/N, bairro Parque Ipanema': '{{empresa_endereco}}',
        'Itaporã': '{{empresa_cidade}}',
        '79890-270': '{{empresa_cep}}',
        'Mateus Leonço Concato': '{{empresa_representante_nome}}',
        'Mateus Leonco Concato': '{{empresa_representante_nome}}',
        '051.536.511-42': '{{empresa_representante_cpf}}',
        '001939969': '{{empresa_representante_rg}}',
        
        # Sistema
        '8,40 kWp': '{{potencia_total}} kWp',
        '8,40': '{{potencia_total}}',
        '10': '{{quantidade_paineis}}',
        
        # Valores
        'doze mil e oitocentos reais': '{{valor_total_extenso}}',
        'R$12.800,00': 'R$ {{valor_total}}',
        '12.800,00': '{{valor_total}}',
        'R$841,45': 'R$ {{valor_parcela}}',
        '841,45': '{{valor_parcela}}',
        'oitocentos e quarenta e um reais e quarenta e cinco centavos': '{{valor_parcela_extenso}}',
        'R$15.146,14': 'R$ {{valor_total_parcelado}}',
        '15.146,14': '{{valor_total_parcelado}}',
        'quinze mil, cento e quarenta e seis reais e quatorze centavos': '{{valor_total_parcelado_extenso}}',
        '18X': '{{numero_parcelas}}X',
        '18x': '{{numero_parcelas}}x',
        
        # Banco
        'Banco Cooperativo Sicredi S.A.': '{{banco_nome}}',
        '0903': '{{banco_agencia}}',
        '50661-1': '{{banco_conta}}',
        '38.068.450 Mateus Leonco Concato': '{{banco_titular}}',
        
        # Prazos
        '45 dias': '{{prazo_execucao_dias}} dias',
        '12 meses': '{{garantia_instalacao_meses}} meses',
        
        # Foro
        'Itaporã/MS': '{{foro_comarca}}',
    }
    
    # Substituir em parágrafos
    for paragrafo in doc.paragraphs:
        for texto_original, chave in substituicoes.items():
            if texto_original in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(texto_original, chave)
    
    # Substituir em tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for texto_original, chave in substituicoes.items():
                    if texto_original in celula.text:
                        celula.text = celula.text.replace(texto_original, chave)
    
    doc.save(output_path)
    print(f"✓ Template criado com chaves: {output_path}")
    print(f"\n📋 Total de chaves substituídas: {len(substituicoes)}")

def main():
    pdf_path = 'CONTRATO - 1 INVERSOR - ORIGINAL.pdf'
    docx_temp = 'contrato_temp.docx'
    docx_template = 'templates/contratos/template_contrato.docx'
    
    print("🔄 Iniciando conversão de PDF para template DOCX...\n")
    
    # Passo 1: Converter PDF para DOCX
    converter_pdf_para_docx(pdf_path, docx_temp)
    
    # Passo 2: Substituir dados por chaves
    substituir_dados_por_chaves(docx_temp, docx_template)
    
    print("\n✅ Conversão concluída!")
    print(f"📄 Template salvo em: {docx_template}")
    print("\n💡 Agora você pode usar este template no sistema")

if __name__ == '__main__':
    main()
