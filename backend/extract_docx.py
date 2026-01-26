from docx import Document

doc = Document('/app/media/templates/PROPOSTA_COMERCIAL_-_MAB.docx')

print('\n=== PARAGRAFOS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'{i}: {p.text}')

print('\n=== TABELAS ===')
for i, table in enumerate(doc.tables):
    print(f'\nTabela {i}:')
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                print(f'  {cell.text}')
