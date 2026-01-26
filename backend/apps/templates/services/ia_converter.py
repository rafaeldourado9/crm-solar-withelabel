from docx import Document
import re

class IATemplateConverterService:
    
    @staticmethod
    def converter_documento(arquivo_path):
        """Converte [chave] para {{CHAVE}}"""
        if not arquivo_path.endswith('.docx'):
            raise ValueError('Apenas arquivos .docx são suportados')
        
        doc = Document(arquivo_path)
        mapeamento = []
        
        # Substituir em parágrafos
        for paragrafo in doc.paragraphs:
            matches = re.findall(r'\[([^\]]+)\]', paragrafo.text)
            for match in matches:
                chave_original = f'[{match}]'
                chave_nova = f'{{{{{match.upper().replace(" ", "_")}}}}}'
                
                for run in paragrafo.runs:
                    if chave_original in run.text:
                        run.text = run.text.replace(chave_original, chave_nova)
                        mapeamento.append((chave_original, chave_nova))
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragrafo in cell.paragraphs:
                        matches = re.findall(r'\[([^\]]+)\]', paragrafo.text)
                        for match in matches:
                            chave_original = f'[{match}]'
                            chave_nova = f'{{{{{match.upper().replace(" ", "_")}}}}}'
                            
                            for run in paragrafo.runs:
                                if chave_original in run.text:
                                    run.text = run.text.replace(chave_original, chave_nova)
                                    mapeamento.append((chave_original, chave_nova))
        
        return doc, mapeamento
