from docx import Document
from io import BytesIO
from decimal import Decimal
from datetime import datetime, timedelta

class TemplateProcessorService:
    
    @staticmethod
    def processar_template(template_path, orcamento, premissa, cliente):
        """Processa template DOCX substituindo variáveis"""
        doc = Document(template_path)
        
        # Preparar dados
        dados = TemplateProcessorService._preparar_dados(orcamento, premissa, cliente)
        
        # Substituir em parágrafos
        for paragrafo in doc.paragraphs:
            TemplateProcessorService._substituir_texto(paragrafo, dados)
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragrafo in cell.paragraphs:
                        TemplateProcessorService._substituir_texto(paragrafo, dados)
        
        # Salvar em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    @staticmethod
    def _substituir_texto(paragrafo, dados):
        """Substitui {{CHAVE}} pelos valores mantendo formatação"""
        texto_original = paragrafo.text
        
        # Verificar se há placeholders
        tem_placeholder = False
        for chave in dados.keys():
            if f'{{{{{chave}}}}}' in texto_original:
                tem_placeholder = True
                break
        
        if not tem_placeholder:
            return
        
        # Substituir mantendo formatação de cada run
        for run in paragrafo.runs:
            for chave, valor in dados.items():
                placeholder = f'{{{{{chave}}}}}'
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(valor))
    
    @staticmethod
    def _preparar_dados(orcamento, premissa, cliente):
        """Prepara dicionário com todos os dados"""
        
        # Calcular valores
        potencia_kwp = (orcamento.potencia_painel * orcamento.quantidade_paineis) / 1000
        hsp = float(premissa.hsp_padrao) if premissa.hsp_padrao else 4.5
        perda = float(premissa.perda_padrao) if premissa.perda_padrao else 20.0
        geracao_mensal = potencia_kwp * hsp * 30 * (1 - perda / 100)
        
        # Data de validade
        data_validade = orcamento.data_criacao + timedelta(days=orcamento.validade_dias)
        
        # Forma de pagamento
        forma_pagamento_texto = 'À vista'
        if orcamento.forma_pagamento != 'avista':
            parcelas = int(orcamento.forma_pagamento)
            forma_pagamento_texto = f'{parcelas}x de R$ {orcamento.valor_parcela:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
        
        dados = {
            # Orçamento
            'NUMERO_ORCAMENTO': orcamento.numero,
            'DATA_ORCAMENTO': orcamento.data_criacao.strftime('%d/%m/%Y'),
            'DATA_CRIACAO': orcamento.data_criacao.strftime('%d/%m/%Y'),  # Alias
            'DATA_VALIDADE': data_validade.strftime('%d/%m/%Y'),
            
            # Cliente
            'NOME_CLIENTE': cliente.nome,
            'CLIENTE_NOME': cliente.nome,  # Alias
            'CPF_CNPJ': cliente.cpf_cnpj or '',
            'TELEFONE': cliente.telefone or '',
            'EMAIL': cliente.email or '',
            'ENDERECO': cliente.endereco or '',
            'CLIENTE_ENDERECO': cliente.endereco or '',  # Alias
            'CIDADE': cliente.cidade or '',
            'CLIENTE_CIDADE': cliente.cidade or '',  # Alias
            'ESTADO': cliente.estado or '',
            'CLIENTE_ESTADO': cliente.estado or '',  # Alias,
            
            # Sistema
            'POTENCIA_KWP': f'{potencia_kwp:.2f}',
            'POTENCIA_TOTAL_KWP': f'{potencia_kwp:.2f}',  # Alias
            'GERACAO_MENSAL': f'{geracao_mensal:.0f}',
            'GERACAO_ANUAL': f'{geracao_mensal * 12:.0f}',
            
            # Equipamentos
            'MARCA_PAINEL': orcamento.marca_painel,
            'PAINEIS_MARCA': orcamento.marca_painel,  # Alias
            'POTENCIA_PAINEL': orcamento.potencia_painel,
            'PAINEIS_POTENCIA': orcamento.potencia_painel,  # Alias
            'QUANTIDADE_PAINEIS': orcamento.quantidade_paineis,
            'PAINEIS_QTD': orcamento.quantidade_paineis,  # Alias
            'MARCA_INVERSOR': orcamento.marca_inversor,
            'INVERSOR_MARCA': orcamento.marca_inversor,  # Alias
            'POTENCIA_INVERSOR': orcamento.potencia_inversor,
            'INVERSOR_POTENCIA': orcamento.potencia_inversor,  # Alias
            'POTENCIA_INVERSOR_KW': f'{orcamento.potencia_inversor / 1000:.1f}',
            'QUANTIDADE_INVERSORES': orcamento.quantidade_inversores,
            'INVERSOR_QTD': orcamento.quantidade_inversores,  # Alias,
            
            # Estrutura
            'TIPO_ESTRUTURA': orcamento.get_tipo_estrutura_display(),
            
            # Valores
            'VALOR_KIT': f'R$ {orcamento.valor_kit:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            'VALOR_ESTRUTURA': f'R$ {orcamento.valor_estrutura:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            'VALOR_MATERIAL_ELETRICO': f'R$ {orcamento.valor_material_eletrico:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            'VALOR_PROJETO': f'R$ {premissa.valor_projeto:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            'VALOR_MONTAGEM': f'R$ {premissa.montagem_por_painel * orcamento.quantidade_paineis:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            'VALOR_TOTAL': f'R$ {orcamento.valor_total:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            'VALOR_FINAL': f'R$ {orcamento.valor_final:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'),
            
            # Pagamento
            'FORMA_PAGAMENTO': forma_pagamento_texto,
            'TAXA_JUROS': f'{orcamento.taxa_juros:.2f}%' if orcamento.taxa_juros > 0 else '0%',
            
            # Vendedor
            'NOME_VENDEDOR': orcamento.vendedor.nome if orcamento.vendedor else '',
            'TELEFONE_VENDEDOR': orcamento.vendedor.telefone if orcamento.vendedor else '',
            'EMAIL_VENDEDOR': orcamento.vendedor.email if orcamento.vendedor else '',
            
            # Premissas
            'HSP': f'{hsp:.2f}',
            'PERDA_SISTEMA': f'{perda:.1f}%',
        }
        
        return dados
