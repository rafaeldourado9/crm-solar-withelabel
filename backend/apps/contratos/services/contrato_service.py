from docx import Document
from io import BytesIO
from ..utils import numero_por_extenso
import os


class ContratoService:
    
    @staticmethod
    def gerar_contrato_docx(contrato):
        """Gera contrato em formato DOCX usando template"""
        
        template_path = '/app/templates/contratos/template_contrato.docx'
        
        # Se não existir template, usar geração padrão
        if not os.path.exists(template_path):
            return ContratoService._gerar_contrato_padrao(contrato)
        
        # Carregar template
        doc = Document(template_path)
        
        cliente = contrato.proposta.orcamento.cliente
        orcamento = contrato.proposta.orcamento
        potencia_kwp = (orcamento.quantidade_paineis * orcamento.potencia_painel) / 1000
        
        # Mapeamento de substituições
        substituicoes = {
            '{{cliente_nome}}': cliente.nome.upper(),
            '{{cliente_cpf_cnpj}}': cliente.cpf_cnpj,
            '{{cliente_rg}}': '',
            '{{cliente_endereco}}': cliente.endereco,
            '{{cliente_bairro}}': cliente.bairro,
            '{{cliente_cidade}}': cliente.cidade,
            '{{cliente_estado}}': cliente.estado,
            '{{cliente_cep}}': cliente.cep,
            '{{empresa_razao_social}}': contrato.empresa_razao_social,
            '{{empresa_cnpj}}': contrato.empresa_cnpj,
            '{{empresa_endereco}}': contrato.empresa_endereco,
            '{{empresa_cidade}}': contrato.empresa_cidade,
            '{{empresa_cep}}': contrato.empresa_cep,
            '{{empresa_representante_nome}}': contrato.empresa_representante_nome,
            '{{empresa_representante_cpf}}': contrato.empresa_representante_cpf,
            '{{empresa_representante_rg}}': contrato.empresa_representante_rg,
            '{{potencia_total}}': f"{potencia_kwp:.2f}",
            '{{quantidade_paineis}}': str(orcamento.quantidade_paineis),
            '{{valor_total}}': f"{contrato.valor_total:,.2f}",
            '{{valor_total_extenso}}': numero_por_extenso(contrato.valor_total),
            '{{valor_parcela}}': f"{contrato.valor_parcela:,.2f}",
            '{{valor_parcela_extenso}}': numero_por_extenso(contrato.valor_parcela),
            '{{numero_parcelas}}': str(contrato.numero_parcelas),
            '{{banco_nome}}': contrato.banco_nome,
            '{{banco_agencia}}': contrato.banco_agencia,
            '{{banco_conta}}': contrato.banco_conta,
            '{{banco_titular}}': contrato.banco_titular,
            '{{prazo_execucao_dias}}': str(contrato.prazo_execucao_dias),
            '{{garantia_instalacao_meses}}': str(contrato.garantia_instalacao_meses),
            '{{foro_comarca}}': contrato.foro_comarca,
        }
        
        # Substituir em parágrafos
        for paragrafo in doc.paragraphs:
            for chave, valor in substituicoes.items():
                if chave in paragrafo.text:
                    for run in paragrafo.runs:
                        if chave in run.text:
                            run.text = run.text.replace(chave, valor)
        
        # Substituir em tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for chave, valor in substituicoes.items():
                            if chave in paragrafo.text:
                                for run in paragrafo.runs:
                                    if chave in run.text:
                                        run.text = run.text.replace(chave, valor)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    @staticmethod
    def _gerar_contrato_padrao(contrato):
        """Fallback: gera contrato padrão se não houver template"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run('CONTRATO PARTICULAR DE COMPRA E INSTALAÇÃO')
        run.bold = True
        run.font.size = Pt(12)
        
        cliente = contrato.proposta.orcamento.cliente
        doc.add_paragraph(f"Contratante: {cliente.nome}")
        doc.add_paragraph(f"Valor: R$ {contrato.valor_total}")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
