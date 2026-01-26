from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PROPOSTA COMERCIAL')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph()

# Informações do Cliente
doc.add_heading('INFORMAÇÕES DO CLIENTE', level=1)
p = doc.add_paragraph()
p.add_run('Nome/Razão Social: ').bold = True
p.add_run('{{CLIENTE_NOME}}')

p = doc.add_paragraph()
p.add_run('CPF/CNPJ: ').bold = True
p.add_run('{{CLIENTE_CPF_CNPJ}}')

p = doc.add_paragraph()
p.add_run('Telefone: ').bold = True
p.add_run('{{CLIENTE_TELEFONE}}')

p = doc.add_paragraph()
p.add_run('E-mail: ').bold = True
p.add_run('{{CLIENTE_EMAIL}}')

p = doc.add_paragraph()
p.add_run('Endereço: ').bold = True
p.add_run('{{CLIENTE_ENDERECO}}, {{CLIENTE_CIDADE}} - {{CLIENTE_ESTADO}}')

doc.add_paragraph()

# Dados da Proposta
doc.add_heading('DADOS DA PROPOSTA', level=1)
p = doc.add_paragraph()
p.add_run('Número: ').bold = True
p.add_run('{{NUMERO_ORCAMENTO}}')

p = doc.add_paragraph()
p.add_run('Data de Criação: ').bold = True
p.add_run('{{DATA_CRIACAO}}')

p = doc.add_paragraph()
p.add_run('Validade: ').bold = True
p.add_run('30 dias')

doc.add_paragraph()

# Dimensionamento
doc.add_heading('DIMENSIONAMENTO DO SISTEMA', level=1)
p = doc.add_paragraph()
p.add_run('Potência Instalada: ').bold = True
p.add_run('{{POTENCIA_TOTAL_KWP}} kWp')

p = doc.add_paragraph()
p.add_run('Geração Estimada: ').bold = True
p.add_run('{{GERACAO_ESTIMADA_KWH}} kWh/mês')

p = doc.add_paragraph()
p.add_run('Área Necessária: ').bold = True
p.add_run('Aproximadamente {{AREA_NECESSARIA}} m²')

doc.add_paragraph()

# Equipamentos
doc.add_heading('EQUIPAMENTOS', level=1)

doc.add_heading('Painéis Solares', level=2)
p = doc.add_paragraph()
p.add_run('Marca: ').bold = True
p.add_run('{{PAINEIS_MARCA}}')

p = doc.add_paragraph()
p.add_run('Modelo: ').bold = True
p.add_run('{{PAINEIS_MODELO}}')

p = doc.add_paragraph()
p.add_run('Potência: ').bold = True
p.add_run('{{PAINEIS_POTENCIA}}')

p = doc.add_paragraph()
p.add_run('Quantidade: ').bold = True
p.add_run('{{PAINEIS_QTD}} unidades')

p = doc.add_paragraph()
p.add_run('Potência Total: ').bold = True
p.add_run('{{PAINEIS_POTENCIA_TOTAL}}')

doc.add_heading('Inversor', level=2)
p = doc.add_paragraph()
p.add_run('Marca: ').bold = True
p.add_run('{{INVERSOR_MARCA}}')

p = doc.add_paragraph()
p.add_run('Modelo: ').bold = True
p.add_run('{{INVERSOR_MODELO}}')

p = doc.add_paragraph()
p.add_run('Potência: ').bold = True
p.add_run('{{INVERSOR_POTENCIA}}')

p = doc.add_paragraph()
p.add_run('Quantidade: ').bold = True
p.add_run('{{INVERSOR_QTD}} unidade(s)')

doc.add_paragraph()

# Tabela de Itens
doc.add_heading('COMPOSIÇÃO DO INVESTIMENTO', level=1)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'

# Cabeçalho
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ITEM'
hdr_cells[1].text = 'QUANTIDADE'
hdr_cells[2].text = 'DESCRIÇÃO'

# Painéis
row = table.rows[1].cells
row[0].text = '1'
row[1].text = '{{PAINEIS_QTD}}'
row[2].text = '{{PAINEIS_MARCA}} - {{PAINEIS_POTENCIA}}'

# Inversor
row = table.rows[2].cells
row[0].text = '2'
row[1].text = '{{INVERSOR_QTD}}'
row[2].text = '{{INVERSOR_MARCA}} - {{INVERSOR_POTENCIA}}'

# Estrutura
row = table.rows[3].cells
row[0].text = '3'
row[1].text = 'Incluso'
row[2].text = 'Estrutura de fixação em alumínio'

# Projeto
row = table.rows[4].cells
row[0].text = '4'
row[1].text = 'Incluso'
row[2].text = 'Projeto, instalação e homologação'

doc.add_paragraph()

# Valores
doc.add_heading('INVESTIMENTO', level=1)
p = doc.add_paragraph()
p.add_run('Valor Total: ').bold = True
run = p.add_run('{{VALOR_FINAL}}')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 128, 0)

p = doc.add_paragraph()
p.add_run('Custo do Kit: ').bold = True
p.add_run('{{VALOR_KIT}}')

p = doc.add_paragraph()
p.add_run('Margem de Lucro: ').bold = True
p.add_run('{{LUCRO_PERCENTUAL}}%')

p = doc.add_paragraph()
p.add_run('Comissão: ').bold = True
p.add_run('{{COMISSAO_PERCENTUAL}}%')

p = doc.add_paragraph()
p.add_run('Impostos: ').bold = True
p.add_run('{{IMPOSTO_PERCENTUAL}}%')

doc.add_paragraph()

# Garantias
doc.add_heading('GARANTIAS', level=1)
p = doc.add_paragraph('✓ Painéis Solares: 25 anos de garantia linear de performance')
p = doc.add_paragraph('✓ Inversor: 10 anos de garantia do fabricante')
p = doc.add_paragraph('✓ Estruturas: 1 ano de garantia')
p = doc.add_paragraph('✓ Mão de Obra: 1 ano de garantia')

doc.add_paragraph()

# Serviços Inclusos
doc.add_heading('SERVIÇOS INCLUSOS', level=1)
p = doc.add_paragraph('✓ Projeto elétrico e estrutural')
p = doc.add_paragraph('✓ Instalação completa do sistema')
p = doc.add_paragraph('✓ Homologação junto à concessionária')
p = doc.add_paragraph('✓ Materiais elétricos (cabos, conectores, proteções)')
p = doc.add_paragraph('✓ Estrutura de fixação')
p = doc.add_paragraph('✓ Treinamento de uso do sistema')

doc.add_paragraph()
doc.add_paragraph()

# Assinaturas
doc.add_paragraph('_' * 50)
p = doc.add_paragraph('{{CLIENTE_NOME}}')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('CLIENTE')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

doc.add_paragraph('_' * 50)
p = doc.add_paragraph('{{VENDEDOR_NOME}}')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('VENDEDOR')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Salvar
doc.save('/app/media/templates/TEMPLATE_PROPOSTA_COMERCIAL.docx')
print('Template criado: /app/media/templates/TEMPLATE_PROPOSTA_COMERCIAL.docx')
