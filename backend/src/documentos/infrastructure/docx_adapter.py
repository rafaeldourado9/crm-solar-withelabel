"""Adapter python-docx: substitui variáveis em templates DOCX."""
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from src.documentos.domain.services import TemplateEngine


def _substituir_paragrafo(paragrafo, variaveis: dict[str, str]) -> None:
    """Substitui variáveis num parágrafo, unindo runs para evitar splits."""
    texto_completo = ''.join(run.text for run in paragrafo.runs)
    if '{{' not in texto_completo:
        return
    texto_novo = TemplateEngine.substituir_texto(texto_completo, variaveis)
    if paragrafo.runs:
        paragrafo.runs[0].text = texto_novo
        for run in paragrafo.runs[1:]:
            run.text = ''


def _substituir_tabelas(doc: Document, variaveis: dict[str, str]) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs:
                    _substituir_paragrafo(paragrafo, variaveis)


def gerar_docx(template_bytes: bytes, variaveis: dict[str, str]) -> bytes:
    """
    Recebe bytes de um template DOCX e um dicionário de variáveis.
    Retorna bytes do DOCX com todas as variáveis substituídas.
    """
    doc = Document(BytesIO(template_bytes))

    for paragrafo in doc.paragraphs:
        _substituir_paragrafo(paragrafo, variaveis)

    _substituir_tabelas(doc, variaveis)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def extrair_variaveis_docx(template_bytes: bytes) -> set[str]:
    """Extrai todas as variáveis {{CHAVE}} presentes no template."""
    doc = Document(BytesIO(template_bytes))
    variaveis: set[str] = set()

    for paragrafo in doc.paragraphs:
        texto = ''.join(run.text for run in paragrafo.runs)
        variaveis |= TemplateEngine.extrair_variaveis(texto)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs:
                    texto = ''.join(run.text for run in paragrafo.runs)
                    variaveis |= TemplateEngine.extrair_variaveis(texto)

    return variaveis
