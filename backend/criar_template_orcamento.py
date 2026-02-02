#!/usr/bin/env python
"""
Script para criar e cadastrar template de orçamento automaticamente
Execute: docker-compose exec backend python criar_template_orcamento.py
"""
import os
import sys
import django

# Configurar Django
sys.path.insert(0, os.path.dirname(__file__))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from docx import Document
from docx.shared import Pt
from django.core.files import File
from apps.templates.models import Template
import io

def criar_template():
    print("=" * 60)
    print("CRIANDO TEMPLATE DE ORÇAMENTO")
    print("=" * 60)
    
    # Verificar se já existe
    templates_existentes = Template.objects.filter(tipo='orcamento', ativo=True)
    if templates_existentes.exists():
        print(f"\n✓ Já existe {templates_existentes.count()} template(s) ativo(s):")
        for t in templates_existentes:
            print(f"  - {t.nome}")
        print("\n✅ Template já cadastrado! Nada a fazer.")
        return
    
    print("\n1. Criando documento DOCX...")
    doc = Document()
    
    # Título
    p = doc.add_paragraph()
    run = p.add_run('ORÇAMENTO - SISTEMA FOTOVOLTAICO')
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = 1  # Center
    
    doc.add_paragraph()
    
    # Orçamento
    p = doc.add_paragraph()
    p.add_run('Orçamento Nº: ').bold = True
    p.add_run('{{NUMERO_ORCAMENTO}}')
    
    p = doc.add_paragraph()
    p.add_run('Data: ').bold = True
    p.add_run('{{DATA_ORCAMENTO}}')
    
    p = doc.add_paragraph()
    p.add_run('Validade: ').bold = True
    p.add_run('{{DATA_VALIDADE}}')
    
    doc.add_paragraph('─' * 60)
    
    # Cliente
    p = doc.add_paragraph()
    run = p.add_run('DADOS DO CLIENTE')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Nome: ').bold = True
    p.add_run('{{NOME_CLIENTE}}')
    
    p = doc.add_paragraph()
    p.add_run('CPF/CNPJ: ').bold = True
    p.add_run('{{CPF_CNPJ}}')
    
    p = doc.add_paragraph()
    p.add_run('Telefone: ').bold = True
    p.add_run('{{TELEFONE}}')
    
    p = doc.add_paragraph()
    p.add_run('Cidade: ').bold = True
    p.add_run('{{CIDADE}} - {{ESTADO}}')
    
    doc.add_paragraph('─' * 60)
    
    # Sistema
    p = doc.add_paragraph()
    run = p.add_run('SISTEMA FOTOVOLTAICO')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Potência: ').bold = True
    p.add_run('{{POTENCIA_KWP}} kWp')
    
    p = doc.add_paragraph()
    p.add_run('Geração Mensal: ').bold = True
    p.add_run('{{GERACAO_MENSAL}} kWh')
    
    p = doc.add_paragraph()
    p.add_run('Geração Anual: ').bold = True
    p.add_run('{{GERACAO_ANUAL}} kWh')
    
    doc.add_paragraph('─' * 60)
    
    # Equipamentos
    p = doc.add_paragraph()
    run = p.add_run('EQUIPAMENTOS')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Painéis: ').bold = True
    p.add_run('{{QUANTIDADE_PAINEIS}}x {{MARCA_PAINEL}} {{POTENCIA_PAINEL}}W')
    
    p = doc.add_paragraph()
    p.add_run('Inversor: ').bold = True
    p.add_run('{{QUANTIDADE_INVERSORES}}x {{MARCA_INVERSOR}} {{POTENCIA_INVERSOR_KW}} kW')
    
    p = doc.add_paragraph()
    p.add_run('Estrutura: ').bold = True
    p.add_run('{{TIPO_ESTRUTURA}}')
    
    doc.add_paragraph('─' * 60)
    
    # Valores
    p = doc.add_paragraph()
    run = p.add_run('VALORES')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Valor do Kit: ').bold = True
    p.add_run('{{VALOR_KIT}}')
    
    p = doc.add_paragraph()
    p.add_run('Estrutura: ').bold = True
    p.add_run('{{VALOR_ESTRUTURA}}')
    
    p = doc.add_paragraph()
    p.add_run('Material Elétrico: ').bold = True
    p.add_run('{{VALOR_MATERIAL_ELETRICO}}')
    
    p = doc.add_paragraph()
    p.add_run('Projeto: ').bold = True
    p.add_run('{{VALOR_PROJETO}}')
    
    p = doc.add_paragraph()
    p.add_run('Montagem: ').bold = True
    p.add_run('{{VALOR_MONTAGEM}}')
    
    doc.add_paragraph('═' * 60)
    
    # Valor Final
    p = doc.add_paragraph()
    p.add_run('VALOR TOTAL: ').bold = True
    run = p.add_run('{{VALOR_TOTAL}}')
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('VALOR FINAL: ').bold = True
    run = p.add_run('{{VALOR_FINAL}}')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('═' * 60)
    
    # Pagamento
    p = doc.add_paragraph()
    run = p.add_run('FORMA DE PAGAMENTO')
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run('{{FORMA_PAGAMENTO}}')
    
    doc.add_paragraph()
    
    # Vendedor
    p = doc.add_paragraph()
    run = p.add_run('VENDEDOR')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('{{NOME_VENDEDOR}}')
    
    p = doc.add_paragraph()
    p.add_run('Tel: {{TELEFONE_VENDEDOR}}')
    
    p = doc.add_paragraph()
    p.add_run('Email: {{EMAIL_VENDEDOR}}')
    
    print("   ✓ Documento criado")
    
    # Salvar em buffer
    print("\n2. Salvando em buffer...")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    print("   ✓ Buffer criado")
    
    # Cadastrar no banco
    print("\n3. Cadastrando no banco de dados...")
    try:
        template = Template.objects.create(
            nome='Orçamento Padrão (Auto)',
            tipo='orcamento',
            arquivo=File(buffer, name='template_orcamento_auto.docx'),
            arquivo_nome='template_orcamento_auto.docx',
            ativo=True
        )
        print(f"   ✓ Template cadastrado! ID: {template.id}")
        print(f"   ✓ Arquivo: {template.arquivo.name}")
    except Exception as e:
        print(f"   ✗ Erro ao cadastrar: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # Testar
    print("\n4. Testando geração...")
    try:
        from apps.orcamentos.models import Orcamento
        from apps.premissas.models import Premissa
        from apps.orcamentos.services.template_processor import TemplateProcessorService
        
        orcamento = Orcamento.objects.first()
        if not orcamento:
            print("   ⚠ Nenhum orçamento para testar")
            print("\n✅ TEMPLATE CADASTRADO COM SUCESSO!")
            print("\nAgora você pode:")
            print("1. Acessar um orçamento no frontend")
            print("2. Clicar em 'Gerar PDF'")
            print("3. O arquivo será baixado automaticamente!")
            return
        
        premissa = Premissa.get_ativa()
        cliente = orcamento.cliente
        
        buffer_test = TemplateProcessorService.processar_template(
            template.arquivo.path,
            orcamento,
            premissa,
            cliente
        )
        
        print(f"   ✓ Teste bem-sucedido! ({len(buffer_test.getvalue())} bytes)")
        
    except Exception as e:
        print(f"   ⚠ Erro no teste: {e}")
        print("   Mas o template foi cadastrado!")
    
    print("\n" + "=" * 60)
    print("✅ SUCESSO! TEMPLATE CADASTRADO")
    print("=" * 60)
    print("\nAgora você pode:")
    print("1. Acessar: http://localhost:5173/orcamentos")
    print("2. Clicar em um orçamento")
    print("3. Clicar em 'Gerar PDF'")
    print("4. O arquivo será baixado automaticamente!")
    print("\nPara personalizar o template:")
    print("1. Acesse: http://localhost:8000/admin/templates/template/")
    print("2. Edite o template e faça upload de um novo arquivo")

if __name__ == '__main__':
    criar_template()
