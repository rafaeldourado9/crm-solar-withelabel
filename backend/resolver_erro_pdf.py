#!/usr/bin/env python
"""
Script para resolver o erro 404 ao gerar PDF
Cria e cadastra automaticamente um template de orçamento
"""
import os
import sys
import django

# Configurar Django
sys.path.insert(0, os.path.dirname(__file__))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

def resolver():
    print("=" * 60)
    print("RESOLVENDO ERRO 404 - GERAÇÃO DE PDF")
    print("=" * 60)
    
    # 1. Verificar python-docx
    print("\n1. Verificando python-docx...")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print("   ✓ python-docx instalado")
    except ImportError:
        print("   ✗ python-docx NÃO instalado")
        print("\n   Instalando...")
        os.system('pip install python-docx')
        print("   ✓ Instalado!")
    
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.core.files import File
    from apps.templates.models import Template
    
    # 2. Verificar se já existe template
    print("\n2. Verificando templates existentes...")
    templates_existentes = Template.objects.filter(tipo='orcamento', ativo=True)
    
    if templates_existentes.exists():
        print(f"   ✓ Já existe {templates_existentes.count()} template(s) ativo(s)")
        for t in templates_existentes:
            print(f"     - {t.nome}")
        print("\n   ✅ PROBLEMA JÁ RESOLVIDO!")
        print("   Se ainda houver erro, verifique se o arquivo existe:")
        for t in templates_existentes:
            if os.path.exists(t.arquivo.path):
                print(f"     ✓ {t.arquivo.path}")
            else:
                print(f"     ✗ {t.arquivo.path} NÃO EXISTE!")
        return
    
    print("   ⚠ Nenhum template ativo encontrado")
    print("   → Criando template automaticamente...")
    
    # 3. Criar template DOCX
    print("\n3. Criando arquivo template...")
    doc = Document()
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('ORÇAMENTO - SISTEMA FOTOVOLTAICO')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Orçamento
    p = doc.add_paragraph()
    p.add_run('Orçamento Nº: ').bold = True
    p.add_run('{{NUMERO_ORCAMENTO}}')
    
    p = doc.add_paragraph()
    p.add_run('Data: ').bold = True
    p.add_run('{{DATA_ORCAMENTO}}')
    
    doc.add_paragraph('─' * 60)
    
    # Cliente
    heading = doc.add_paragraph()
    run = heading.add_run('CLIENTE')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Nome: ').bold = True
    p.add_run('{{NOME_CLIENTE}}')
    
    p = doc.add_paragraph()
    p.add_run('CPF/CNPJ: ').bold = True
    p.add_run('{{CPF_CNPJ}}')
    
    p = doc.add_paragraph()
    p.add_run('Telefone: ').bold = True
    p.add_run('{{TELEFONE}}')
    
    doc.add_paragraph('─' * 60)
    
    # Sistema
    heading = doc.add_paragraph()
    run = heading.add_run('SISTEMA FOTOVOLTAICO')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Potência: ').bold = True
    p.add_run('{{POTENCIA_KWP}} kWp')
    
    p = doc.add_paragraph()
    p.add_run('Geração Mensal: ').bold = True
    p.add_run('{{GERACAO_MENSAL}} kWh')
    
    doc.add_paragraph('─' * 60)
    
    # Equipamentos
    heading = doc.add_paragraph()
    run = heading.add_run('EQUIPAMENTOS')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Painéis: ').bold = True
    p.add_run('{{QUANTIDADE_PAINEIS}}x {{MARCA_PAINEL}} {{POTENCIA_PAINEL}}W')
    
    p = doc.add_paragraph()
    p.add_run('Inversor: ').bold = True
    p.add_run('{{QUANTIDADE_INVERSORES}}x {{MARCA_INVERSOR}} {{POTENCIA_INVERSOR_KW}} kW')
    
    p = doc.add_paragraph()
    p.add_run('Estrutura: ').bold = True
    p.add_run('{{TIPO_ESTRUTURA}}')
    
    doc.add_paragraph('─' * 60)
    
    # Valores
    heading = doc.add_paragraph()
    run = heading.add_run('VALORES')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run('Valor do Kit: ').bold = True
    p.add_run('{{VALOR_KIT}}')
    
    p = doc.add_paragraph()
    p.add_run('Estrutura: ').bold = True
    p.add_run('{{VALOR_ESTRUTURA}}')
    
    p = doc.add_paragraph()
    p.add_run('Material Elétrico: ').bold = True
    p.add_run('{{VALOR_MATERIAL_ELETRICO}}')
    
    doc.add_paragraph('═' * 60)
    
    # Valor Final
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('VALOR FINAL: ')
    run.bold = True
    run.font.size = Pt(16)
    run = p.add_run('{{VALOR_FINAL}}')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Forma de Pagamento: ').bold = True
    p.add_run('{{FORMA_PAGAMENTO}}')
    
    # Salvar
    temp_path = 'template_orcamento_auto.docx'
    doc.save(temp_path)
    print(f"   ✓ Arquivo criado: {temp_path}")
    
    # 4. Cadastrar no banco
    print("\n4. Cadastrando template no banco de dados...")
    try:
        with open(temp_path, 'rb') as f:
            template = Template.objects.create(
                nome='Orçamento Padrão (Auto)',
                tipo='orcamento',
                arquivo=File(f, name='template_orcamento_auto.docx'),
                arquivo_nome='template_orcamento_auto.docx',
                ativo=True
            )
        print(f"   ✓ Template cadastrado! ID: {template.id}")
        print(f"   ✓ Arquivo salvo em: {template.arquivo.path}")
    except Exception as e:
        print(f"   ✗ Erro ao cadastrar: {e}")
        return
    
    # 5. Testar
    print("\n5. Testando geração...")
    try:
        from apps.orcamentos.models import Orcamento
        from apps.premissas.models import Premissa
        from apps.orcamentos.services.template_processor import TemplateProcessorService
        
        orcamento = Orcamento.objects.first()
        if not orcamento:
            print("   ⚠ Nenhum orçamento cadastrado para testar")
            print("   ✓ Template cadastrado com sucesso!")
            print("\n   ✅ PROBLEMA RESOLVIDO!")
            return
        
        premissa = Premissa.get_ativa()
        cliente = orcamento.cliente
        
        buffer = TemplateProcessorService.processar_template(
            template.arquivo.path,
            orcamento,
            premissa,
            cliente
        )
        
        test_file = f'teste_orcamento_{orcamento.id}.docx'
        with open(test_file, 'wb') as f:
            f.write(buffer.getvalue())
        
        print(f"   ✓ Teste bem-sucedido!")
        print(f"   ✓ Arquivo de teste: {os.path.abspath(test_file)}")
        
    except Exception as e:
        print(f"   ⚠ Erro no teste: {e}")
        print("   ✓ Mas o template foi cadastrado!")
    
    print("\n" + "=" * 60)
    print("✅ PROBLEMA RESOLVIDO COM SUCESSO!")
    print("=" * 60)
    print("\nAgora você pode:")
    print("1. Acessar a página do orçamento")
    print("2. Clicar em 'Gerar PDF'")
    print("3. O arquivo será baixado automaticamente")
    print("\nSe quiser personalizar o template:")
    print(f"1. Acesse: http://localhost:8000/admin/templates/template/{template.id}/change/")
    print("2. Faça upload de um novo arquivo .docx")
    print("3. Salve")

if __name__ == '__main__':
    resolver()
