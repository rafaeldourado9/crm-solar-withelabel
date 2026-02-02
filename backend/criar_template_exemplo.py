#!/usr/bin/env python
"""
Script para criar template de orçamento de exemplo
"""
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ ERRO: python-docx não instalado!")
    print("\nInstale com: pip install python-docx")
    sys.exit(1)

def criar_template():
    print("Criando template de orçamento...")
    
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('ORÇAMENTO - SISTEMA FOTOVOLTAICO')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    # Informações do Orçamento
    p = doc.add_paragraph()
    p.add_run('Orçamento Nº: ').bold = True
    p.add_run('{{NUMERO_ORCAMENTO}}')
    
    p = doc.add_paragraph()
    p.add_run('Data: ').bold = True
    p.add_run('{{DATA_ORCAMENTO}}')
    
    p = doc.add_paragraph()
    p.add_run('Validade: ').bold = True
    p.add_run('{{DATA_VALIDADE}}')
    
    doc.add_paragraph('─' * 80)
    
    # Cliente
    heading = doc.add_paragraph()
    run = heading.add_run('DADOS DO CLIENTE')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    p = doc.add_paragraph()
    p.add_run('Nome: ').bold = True
    p.add_run('{{NOME_CLIENTE}}')
    
    p = doc.add_paragraph()
    p.add_run('CPF/CNPJ: ').bold = True
    p.add_run('{{CPF_CNPJ}}')
    
    p = doc.add_paragraph()
    p.add_run('Telefone: ').bold = True
    p.add_run('{{TELEFONE}}')
    
    p = doc.add_paragraph()
    p.add_run('E-mail: ').bold = True
    p.add_run('{{EMAIL}}')
    
    p = doc.add_paragraph()
    p.add_run('Cidade: ').bold = True
    p.add_run('{{CIDADE}} - {{ESTADO}}')
    
    doc.add_paragraph('─' * 80)
    
    # Sistema
    heading = doc.add_paragraph()
    run = heading.add_run('SISTEMA FOTOVOLTAICO')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    p = doc.add_paragraph()
    p.add_run('Potência do Sistema: ').bold = True
    run = p.add_run('{{POTENCIA_KWP}} kWp')
    run.font.color.rgb = RGBColor(0, 153, 0)
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run('Geração Mensal Estimada: ').bold = True
    p.add_run('{{GERACAO_MENSAL}} kWh')
    
    p = doc.add_paragraph()
    p.add_run('Geração Anual Estimada: ').bold = True
    p.add_run('{{GERACAO_ANUAL}} kWh')
    
    doc.add_paragraph('─' * 80)
    
    # Equipamentos
    heading = doc.add_paragraph()
    run = heading.add_run('EQUIPAMENTOS')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    p = doc.add_paragraph()
    p.add_run('Painéis Solares: ').bold = True
    p.add_run('{{QUANTIDADE_PAINEIS}}x {{MARCA_PAINEL}} {{POTENCIA_PAINEL}}W')
    
    p = doc.add_paragraph()
    p.add_run('Inversor: ').bold = True
    p.add_run('{{QUANTIDADE_INVERSORES}}x {{MARCA_INVERSOR}} {{POTENCIA_INVERSOR_KW}} kW')
    
    p = doc.add_paragraph()
    p.add_run('Estrutura: ').bold = True
    p.add_run('{{TIPO_ESTRUTURA}}')
    
    doc.add_paragraph('─' * 80)
    
    # Valores
    heading = doc.add_paragraph()
    run = heading.add_run('COMPOSIÇÃO DE VALORES')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Tabela de valores
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    items = [
        ('Valor do Kit', '{{VALOR_KIT}}'),
        ('Estrutura', '{{VALOR_ESTRUTURA}}'),
        ('Material Elétrico', '{{VALOR_MATERIAL_ELETRICO}}'),
        ('Projeto', '{{VALOR_PROJETO}}'),
        ('Montagem', '{{VALOR_MONTAGEM}}'),
    ]
    
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Valor Total
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('VALOR TOTAL: ')
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run('{{VALOR_TOTAL}}')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Valor Final
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('VALOR FINAL: ')
    run.bold = True
    run.font.size = Pt(16)
    run = p.add_run('{{VALOR_FINAL}}')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 153, 0)
    
    doc.add_paragraph('═' * 80)
    
    # Pagamento
    heading = doc.add_paragraph()
    run = heading.add_run('FORMA DE PAGAMENTO')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    p = doc.add_paragraph()
    run = p.add_run('{{FORMA_PAGAMENTO}}')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Vendedor
    heading = doc.add_paragraph()
    run = heading.add_run('INFORMAÇÕES DO VENDEDOR')
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run('{{NOME_VENDEDOR}}')
    
    p = doc.add_paragraph()
    p.add_run('Telefone: {{TELEFONE_VENDEDOR}}')
    
    p = doc.add_paragraph()
    p.add_run('E-mail: {{EMAIL_VENDEDOR}}')
    
    # Salvar
    output_path = 'template_orcamento_exemplo.docx'
    doc.save(output_path)
    
    print(f"\n✅ Template criado com sucesso!")
    print(f"📄 Arquivo: {os.path.abspath(output_path)}")
    print(f"\n📋 Próximos passos:")
    print(f"1. Acesse: http://localhost:8000/admin/templates/template/")
    print(f"2. Clique em 'Adicionar Template'")
    print(f"3. Faça upload do arquivo: {output_path}")
    print(f"4. Marque como 'Ativo'")
    print(f"5. Salve")

if __name__ == '__main__':
    criar_template()
