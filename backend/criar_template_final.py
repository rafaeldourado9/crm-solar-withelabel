from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PROPOSTA\nCOMERCIAL')
run.font.size = Pt(24)
run.font.bold = True

p = doc.add_paragraph()
p.add_run('Cliente:\n').bold = True
p.add_run('{{CLIENTE_NOME}}')

p = doc.add_paragraph()
p.add_run('Potência:\n').bold = True
p.add_run('{{POTENCIA_TOTAL_KWP}} kWp')

doc.add_page_break()

# PROPOSTA
doc.add_heading('PROPOSTA', level=1)
doc.add_heading('INFORMAÇÕES DO CLIENTE', level=2)

p = doc.add_paragraph()
p.add_run('Nome/Razão Social: ').bold = True
p.add_run('{{CLIENTE_NOME}}')

p = doc.add_paragraph()
p.add_run('Endereço: ').bold = True
p.add_run('{{CLIENTE_ENDERECO}}')

p = doc.add_paragraph()
p.add_run('Bairro: ').bold = True
p.add_run('{{CLIENTE_BAIRRO}}')
p.add_run('  Cidade: ').bold = True
p.add_run('{{CLIENTE_CIDADE}}')

p = doc.add_paragraph()
p.add_run('Estado: ').bold = True
p.add_run('{{CLIENTE_ESTADO}}')

doc.add_paragraph()

# DADOS ADICIONAIS
doc.add_heading('DADOS ADICIONAIS', level=2)
p = doc.add_paragraph()
p.add_run('Economia Esperada (%): ').bold = True
p.add_run('{{ECONOMIA_PERCENTUAL}}%')

p = doc.add_paragraph()
p.add_run('Data de Criação: ').bold = True
p.add_run('{{DATA_CRIACAO}}')
p.add_run('  Validade da Proposta: ').bold = True
p.add_run('{{VALIDADE_PROPOSTA}}')

doc.add_paragraph()

# DIMENSIONAMENTO
doc.add_heading('DIMENSIONAMENTO - SISTEMA GERADOR ON GRID', level=2)
p = doc.add_paragraph()
p.add_run('Geração de Energia Requerida: ').bold = True
p.add_run('{{GERACAO_ESTIMADA_KWH}} kWh/mês')

p = doc.add_paragraph()
p.add_run('Potência Instalada: ').bold = True
p.add_run('{{POTENCIA_TOTAL_KWP}} kWp')

p = doc.add_paragraph()
p.add_run('Potência das placas: ').bold = True
p.add_run('{{PAINEIS_POTENCIA}}')

p = doc.add_paragraph()
p.add_run('Quantidade de placas: ').bold = True
p.add_run('{{PAINEIS_QTD}}')

p = doc.add_paragraph()
p.add_run('Área necessária: ').bold = True
p.add_run('{{AREA_NECESSARIA}} m²')

doc.add_paragraph()

# EQUIPAMENTO
doc.add_heading('EQUIPAMENTO', level=2)

doc.add_heading('PLACA SOLAR', level=3)
p = doc.add_paragraph()
p.add_run('MÓDULO {{PAINEIS_MARCA}} {{PAINEIS_POTENCIA}}')

doc.add_heading('INVERSOR', level=3)
p = doc.add_paragraph()
p.add_run('INVERSOR {{INVERSOR_MARCA}} {{INVERSOR_POTENCIA}}')

doc.add_paragraph()

# TABELA
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ITEM'
hdr_cells[1].text = 'QUANTIDADE'

row = table.rows[1].cells
row[0].text = 'MÓDULO {{PAINEIS_MARCA}} {{PAINEIS_POTENCIA}}'
row[1].text = '{{PAINEIS_QTD}}'

row = table.rows[2].cells
row[0].text = 'INVERSOR {{INVERSOR_MARCA}} {{INVERSOR_POTENCIA}}'
row[1].text = '{{INVERSOR_QTD}}'

row = table.rows[3].cells
row[0].text = 'PROJETO E INSTALAÇÃO'
row[1].text = 'Incluso'

p = doc.add_paragraph()
p.add_run('VALOR TOTAL: R$ {{VALOR_FINAL}}').bold = True

doc.add_paragraph()

# GARANTIA
doc.add_heading('GARANTIA', level=2)
doc.add_paragraph('Placas: 25 anos (Linear)')
doc.add_paragraph('Inversor: 10 anos')
doc.add_paragraph('Estruturas: 1 ano')
doc.add_paragraph('Mão de Obra: {{GARANTIA_SERVICO}}')

doc.add_paragraph()

# DETALHES E INVESTIMENTO
doc.add_heading('DETALHES E INVESTIMENTO', level=2)
doc.add_paragraph('Outros itens inclusos...')

doc.add_paragraph('✓ Estrutura em alumínio ou galvanizado a fogo para fixação (perfis e suportes)')
doc.add_paragraph('✓ Materiais elétricos')
doc.add_paragraph('✓ Projeto de regularização junto à concessionária')
doc.add_paragraph('✓ Instalação do sistema solar fotovoltaico')

doc.add_paragraph()

# ANÁLISE FINANCEIRA
doc.add_heading('ANÁLISE FINANCEIRA', level=2)
p = doc.add_paragraph()
p.add_run('Valor da Proposta: R$ {{VALOR_FINAL}}').bold = True

p = doc.add_paragraph()
p.add_run('Tempo de vida mínima: ').bold = True
p.add_run('{{TEMPO_VIDA_SISTEMA}}')

doc.add_paragraph()

# INFORMAÇÕES DA EMPRESA
doc.add_heading('INFORMAÇÕES DA EMPRESA:', level=2)
p = doc.add_paragraph()
p.add_run('Razão Social: ').bold = True
p.add_run('{{EMPRESA_RAZAO_SOCIAL}}')

p = doc.add_paragraph()
p.add_run('CNPJ: ').bold = True
p.add_run('{{EMPRESA_CNPJ}}')

p = doc.add_paragraph()
p.add_run('Telefone: ').bold = True
p.add_run('{{EMPRESA_TELEFONE}}')

p = doc.add_paragraph()
p.add_run('E-mail: ').bold = True
p.add_run('{{EMPRESA_EMAIL}}')

doc.add_paragraph()

# DADOS BANCÁRIOS
doc.add_heading('DADOS BANCÁRIOS', level=2)
p = doc.add_paragraph()
p.add_run('Banco: ').bold = True
p.add_run('{{BANCO_NOME}}')

p = doc.add_paragraph()
p.add_run('Agência: ').bold = True
p.add_run('{{BANCO_AGENCIA}}')

p = doc.add_paragraph()
p.add_run('Conta: ').bold = True
p.add_run('{{BANCO_CONTA}}')

p = doc.add_paragraph()
p.add_run('Razão Social: ').bold = True
p.add_run('{{EMPRESA_RAZAO_SOCIAL}}')

p = doc.add_paragraph()
p.add_run('CNPJ: ').bold = True
p.add_run('{{EMPRESA_CNPJ}}')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# ASSINATURAS
doc.add_paragraph('_' * 60)
p = doc.add_paragraph('{{CLIENTE_NOME}}')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('CLIENTE')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

doc.add_paragraph('_' * 60)
p = doc.add_paragraph('{{VENDEDOR_NOME}} - {{EMPRESA_CNPJ}}')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('VENDEDOR')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('/app/media/templates/PROPOSTA_COMERCIAL_TEMPLATE.docx')
print('✓ Template criado com sucesso!')
print('\nChaves disponíveis:')
print('- {{CLIENTE_NOME}}, {{CLIENTE_ENDERECO}}, {{CLIENTE_BAIRRO}}, {{CLIENTE_CIDADE}}, {{CLIENTE_ESTADO}}')
print('- {{POTENCIA_TOTAL_KWP}}, {{GERACAO_ESTIMADA_KWH}}, {{AREA_NECESSARIA}}')
print('- {{PAINEIS_MARCA}}, {{PAINEIS_POTENCIA}}, {{PAINEIS_QTD}}')
print('- {{INVERSOR_MARCA}}, {{INVERSOR_POTENCIA}}, {{INVERSOR_QTD}}')
print('- {{VALOR_FINAL}}, {{ECONOMIA_PERCENTUAL}}, {{DATA_CRIACAO}}, {{VALIDADE_PROPOSTA}}')
print('- {{GARANTIA_SERVICO}}, {{TEMPO_VIDA_SISTEMA}}')
print('- {{EMPRESA_RAZAO_SOCIAL}}, {{EMPRESA_CNPJ}}, {{EMPRESA_TELEFONE}}, {{EMPRESA_EMAIL}}')
print('- {{BANCO_NOME}}, {{BANCO_AGENCIA}}, {{BANCO_CONTA}}')
print('- {{VENDEDOR_NOME}}')
